# backend/app/services/maintenance_report_service.py

import os
import io
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from openai import OpenAI
from sqlalchemy.orm import Session
from ..models import MaintenanceExecution, MaintenanceProtocol, Machine, Organization, User


class MaintenanceReportService:
    """Service for generating maintenance execution reports with AI insights"""
    
    def __init__(self, db: Session):
        self.db = db
        # Get OpenAI API key from environment and initialize client
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_client = None
        if self.openai_api_key:
            self.openai_client = OpenAI(api_key=self.openai_api_key)
    
    def get_execution_data(self, execution_id: str) -> Dict[str, Any]:
        """Fetch all data needed for the report"""
        execution = self.db.query(MaintenanceExecution).filter(
            MaintenanceExecution.id == execution_id
        ).first()
        
        if not execution:
            raise ValueError(f"Execution {execution_id} not found")
        
        # Get related data
        machine = execution.machine
        protocol = execution.protocol
        organization = machine.customer_organization if machine else None
        performed_by = execution.performed_by
        
        # Get checklist completions with items
        checklist_data = []
        for completion in execution.checklist_completions:
            item = completion.checklist_item
            checklist_data.append({
                'item_id': str(item.id),
                'description': item.item_description,
                'category': item.item_category or 'General',
                'is_completed': completion.is_completed,
                'notes': completion.notes or '',
                'completed_at': completion.completed_at,
                'part_usage': completion.part_usage
            })
        
        return {
            'execution': execution,
            'machine': machine,
            'protocol': protocol,
            'organization': organization,
            'performed_by': performed_by,
            'checklist_data': checklist_data
        }
    
    def generate_ai_insights(self, checklist_item: Dict[str, Any]) -> str:
        """Generate AI insights for a checklist item"""
        if not self.openai_client:
            return "AI insights unavailable (API key not configured)"
        
        try:
            # Build a detailed prompt that emphasizes the specific task
            prompt = f"""Analyze this specific AutoBoss maintenance task and provide targeted professional insights (2-3 sentences):

TASK: "{checklist_item['description']}"
Category: {checklist_item['category']}
Status: {'Completed' if checklist_item['is_completed'] else 'Not Completed'}
Technician Notes: {checklist_item['notes'] if checklist_item['notes'] else 'None provided'}

INSTRUCTIONS:
1. Read the task description carefully and identify what specific component or system it addresses
2. Provide insights ONLY about that specific component/system mentioned in the task
3. Explain why THIS SPECIFIC task matters for AutoBoss operation
4. Describe potential problems if this specific task is not done properly
5. Give actionable recommendations for this specific task

IMPORTANT: 
- Focus ONLY on what the task description mentions
- Do NOT discuss walking wheels unless the task specifically mentions them
- Do NOT discuss net positioning unless the task specifically mentions it
- Do NOT provide generic advice - be specific to this task
- If the task is about electronics, focus on electronics
- If the task is about hydraulics, focus on hydraulics
- If the task is about cleaning brushes, focus on brushes

AutoBoss context (use only if relevant to the task): Automated net cleaning machine for aquaculture, operates in marine environments with saltwater exposure."""

            response = self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an AutoBoss maintenance expert. Analyze each task individually and provide insights ONLY about the specific component or system mentioned in that task. Do not default to discussing walking wheels or net positioning unless the task explicitly mentions them. Be precise and task-specific."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=200,
                temperature=0.7
            )
            
            return response.choices[0].message.content.strip()
        except Exception as e:
            return f"AI insight generation failed: {str(e)}"
    
    async def generate_docx_report(self, execution_id: str) -> io.BytesIO:
        """Generate a DOCX report for the maintenance execution"""
        data = self.get_execution_data(execution_id)
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Maintenance Execution Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add execution metadata
        doc.add_heading('Execution Details', level=1)
        
        metadata_table = doc.add_table(rows=7, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata_rows = [
            ('Organization:', data['organization'].name if data['organization'] else 'N/A'),
            ('Machine:', f"{data['machine'].name} ({data['machine'].serial_number})" if data['machine'] else 'N/A'),
            ('Protocol:', data['protocol'].name if data['protocol'] else 'Custom Maintenance'),
            ('Performed By:', data['performed_by'].username if data['performed_by'] else 'N/A'),
            ('Date:', data['execution'].performed_date.strftime('%Y-%m-%d %H:%M') if data['execution'].performed_date else 'N/A'),
            ('Machine Hours:', str(data['execution'].machine_hours_at_service) if data['execution'].machine_hours_at_service else 'N/A'),
            ('Status:', data['execution'].status.value if data['execution'].status else 'N/A')
        ]
        
        for idx, (label, value) in enumerate(metadata_rows):
            metadata_table.rows[idx].cells[0].text = label
            metadata_table.rows[idx].cells[1].text = value
            # Bold the labels
            metadata_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Add execution notes if any
        if data['execution'].notes:
            doc.add_heading('General Notes', level=1)
            doc.add_paragraph(data['execution'].notes)
            doc.add_paragraph()
        
        # Add checklist items
        doc.add_heading('Maintenance Checklist', level=1)
        
        for idx, item in enumerate(data['checklist_data'], 1):
            # Item header
            item_heading = doc.add_heading(f"Task {idx}: {item['description']}", level=2)
            
            # Item details
            details = doc.add_paragraph()
            details.add_run('Category: ').bold = True
            details.add_run(f"{item['category']}\n")
            details.add_run('Status: ').bold = True
            status_run = details.add_run(f"{'✓ Completed' if item['is_completed'] else '✗ Not Completed'}\n")
            status_run.font.color.rgb = RGBColor(0, 128, 0) if item['is_completed'] else RGBColor(255, 0, 0)
            
            if item['completed_at']:
                details.add_run('Completed At: ').bold = True
                details.add_run(f"{item['completed_at'].strftime('%Y-%m-%d %H:%M')}\n")
            
            # Technician notes
            if item['notes']:
                doc.add_heading('Technician Notes:', level=3)
                doc.add_paragraph(item['notes'], style='Intense Quote')
            
            # AI-generated insights
            doc.add_heading('Expert Insights:', level=3)
            ai_insight = self.generate_ai_insights(item)
            insight_para = doc.add_paragraph(ai_insight, style='Intense Quote')
            insight_para.paragraph_format.left_indent = Inches(0.5)
            
            doc.add_paragraph()  # Spacing
        
        # Add footer
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.add_run(f'Report Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n').italic = True
        footer.add_run('ABParts Maintenance Management System').italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    async def generate_pdf_report(self, execution_id: str) -> io.BytesIO:
        """Generate a PDF report for the maintenance execution"""
        data = self.get_execution_data(execution_id)
        
        # Create PDF buffer
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#374151'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Add title
        elements.append(Paragraph('Maintenance Execution Report', title_style))
        elements.append(Spacer(1, 12))
        
        # Add metadata table
        metadata_data = [
            ['Organization:', data['organization'].name if data['organization'] else 'N/A'],
            ['Machine:', f"{data['machine'].name} ({data['machine'].serial_number})" if data['machine'] else 'N/A'],
            ['Protocol:', data['protocol'].name if data['protocol'] else 'Custom Maintenance'],
            ['Performed By:', data['performed_by'].username if data['performed_by'] else 'N/A'],
            ['Date:', data['execution'].performed_date.strftime('%Y-%m-%d %H:%M') if data['execution'].performed_date else 'N/A'],
            ['Machine Hours:', str(data['execution'].machine_hours_at_service) if data['execution'].machine_hours_at_service else 'N/A'],
            ['Status:', data['execution'].status.value if data['execution'].status else 'N/A']
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        
        elements.append(metadata_table)
        elements.append(Spacer(1, 20))
        
        # Add general notes if any
        if data['execution'].notes:
            elements.append(Paragraph('General Notes', heading_style))
            elements.append(Paragraph(data['execution'].notes, styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # Add checklist items
        elements.append(Paragraph('Maintenance Checklist', heading_style))
        elements.append(Spacer(1, 12))
        
        for idx, item in enumerate(data['checklist_data'], 1):
            # Task header
            task_style = ParagraphStyle(
                'TaskHeader',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.HexColor('#1f2937'),
                spaceAfter=6
            )
            elements.append(Paragraph(f"Task {idx}: {item['description']}", task_style))
            
            # Task details
            status_color = 'green' if item['is_completed'] else 'red'
            status_text = '✓ Completed' if item['is_completed'] else '✗ Not Completed'
            
            task_details = f"""
            <b>Category:</b> {item['category']}<br/>
            <b>Status:</b> <font color="{status_color}">{status_text}</font><br/>
            """
            
            if item['completed_at']:
                task_details += f"<b>Completed At:</b> {item['completed_at'].strftime('%Y-%m-%d %H:%M')}<br/>"
            
            elements.append(Paragraph(task_details, styles['Normal']))
            elements.append(Spacer(1, 6))
            
            # Technician notes
            if item['notes']:
                elements.append(Paragraph('<b>Technician Notes:</b>', styles['Normal']))
                note_style = ParagraphStyle(
                    'NoteStyle',
                    parent=styles['Normal'],
                    leftIndent=20,
                    textColor=colors.HexColor('#4b5563'),
                    fontSize=9
                )
                elements.append(Paragraph(item['notes'], note_style))
                elements.append(Spacer(1, 6))
            
            # AI insights
            elements.append(Paragraph('<b>Expert Insights:</b>', styles['Normal']))
            ai_insight = self.generate_ai_insights(item)
            insight_style = ParagraphStyle(
                'InsightStyle',
                parent=styles['Normal'],
                leftIndent=20,
                textColor=colors.HexColor('#6366f1'),
                fontSize=9,
                fontName='Helvetica-Oblique'
            )
            elements.append(Paragraph(ai_insight, insight_style))
            elements.append(Spacer(1, 12))
        
        # Add footer
        elements.append(PageBreak())
        footer_text = f"""
        <para align=center>
        <i>Report Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</i><br/>
        <i>ABParts Maintenance Management System</i>
        </para>
        """
        elements.append(Paragraph(footer_text, styles['Normal']))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        
        return buffer
